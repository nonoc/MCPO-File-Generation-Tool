"""
Word (DOCX) exporter module using python-docx.
"""

import os
import logging
from typing import Any, List, Optional

from docx import Document
from docx.shared import Pt as DocxPt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from io import BytesIO

# Import from new modules (try relative first)
try:
    from ..shared.utils import _generate_unique_folder, _public_url, render_text_with_emojis
    from ..processing.text_utils import _normalize_markup_text, _parse_paragraph_segments
    from ..processing.block_utils import _finalize_normalized_block
except ImportError:
    from shared.utils import _generate_unique_folder, _public_url, render_text_with_emojis
    from processing.text_utils import _normalize_markup_text, _parse_paragraph_segments
    from processing.block_utils import _finalize_normalized_block

log = logging.getLogger(__name__)


def _filter_title_blocks_from_content(content, has_root_title=False):
    """
    Filter out title blocks from content when a root-level title is provided.
    
    Args:
        content: Content to filter (string, dict, or list)
        has_root_title: If True, filter out title blocks; otherwise return content unchanged
        
    Returns:
        Filtered content with title blocks removed
    """
    if not has_root_title:
        return content
    
    if isinstance(content, str):
        # For markdown strings, remove the first # heading if it exists
        lines = content.splitlines()
        if lines and lines[0].strip().startswith("# "):
            return "\n".join(lines[1:]).lstrip()
        return content
    
    if isinstance(content, dict):
        content_type = (content.get("type") or "").strip().lower()
        if content_type == "title":
            return []
        # If it's a dict that's not a title, wrap in list
        return [content]
    
    if isinstance(content, list):
        # Filter out title blocks while preserving other content
        filtered = []
        for item in content:
            if isinstance(item, dict):
                item_type = (item.get("type") or "").strip().lower()
                if item_type != "title":
                    filtered.append(item)
            elif isinstance(item, str):
                if not item.strip().startswith("# "):
                    filtered.append(item)
            else:
                filtered.append(item)
        return filtered if filtered else []
    
    return content


def _create_word(content: Any, filename: str, folder_path: Optional[str] = None, title: Optional[str] = None) -> dict:
    """
    Create a Word document from content.
    
    Args:
        content: Content to render (string, dict, or list)
        filename: Output filename
        folder_path: Optional folder path for output
        title: Optional document title
        
    Returns:
        Dict with url and path to generated file
    """
    log.debug("Creating Word document")

    if isinstance(content, str):
        from ..processing.block_utils import _convert_markdown_to_structured
        content = _convert_markdown_to_structured(content)
    elif isinstance(content, dict):
        content = [content]
    elif not isinstance(content, list):
        content = []

    # Block-first title logic: title blocks take precedence over root-level title
    # Step 1: Check normalized content for type="title" block first
    normalized_for_title = _normalize_content_for_export(content)
    title_from_block = None
    for block in normalized_for_title:
        if isinstance(block, dict) and (block.get("type") or "").strip().lower() == "title":
            title_from_block = block.get("text")
            break
    
    # Step 2: Use title from block if found, otherwise use root-level title
    if title_from_block:
        title = title_from_block
        # Filter title blocks from content to avoid duplication (has_root_title=True to trigger filtering)
        content = _filter_title_blocks_from_content(content, has_root_title=True)
    else:
        # No title block found - use root-level title if provided
        has_root_title = bool(title)
        # Filter title blocks to avoid duplication when root title exists
        if has_root_title:
            content = _filter_title_blocks_from_content(content, has_root_title=True)

    if folder_path is None:
        folder_path = _generate_unique_folder()
        
    if filename:
        filepath = os.path.join(folder_path, filename)
        os.makedirs(os.path.dirname(filepath), exist_ok=True)
        fname = filename
    else:
        from ..shared.utils import _generate_filename
        filepath, fname = _generate_filename(folder_path, "docx")

    # Handle template
    from ..shared.constants import DOCX_TEMPLATE, DOCX_TEMPLATE_PATH
    use_template = False
    doc = None

    if DOCX_TEMPLATE:
        try:
            src = DOCX_TEMPLATE
            if hasattr(DOCX_TEMPLATE, "paragraphs") and hasattr(DOCX_TEMPLATE, "save"):
                buf = BytesIO()
                DOCX_TEMPLATE.save(buf)
                buf.seek(0)
                src = buf

            doc = Document(src)
            use_template = True
            log.debug("Using DOCX template")

            for element in doc.element.body:
                if element.tag.endswith('}p') or element.tag.endswith('}tbl'):
                    doc.element.body.remove(element)

        except Exception as e:
            log.warning(f"Failed to load DOCX template: {e}")
            use_template = False
            doc = None

    if not use_template:
        doc = Document()
        log.debug("Creating new Word document without template")

    # Add title if provided
    if title:
        title_paragraph = doc.add_paragraph(title)
        try:
            title_paragraph.style = doc.styles['Title']
        except KeyError:
            try:
                title_paragraph.style = doc.styles['Heading 1']
            except KeyError:
                run = title_paragraph.runs[0] if title_paragraph.runs else title_paragraph.add_run()
                run.font.size = DocxPt(20)
                run.font.bold = True
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        log.debug("Document title added")

    # Process content
    normalized_content = _normalize_content_for_export(content)
    if normalized_content:
        _render_structured_docx(doc, normalized_content)
    else:
        _add_docx_paragraph(doc, "")

    doc.save(filepath)
    return {"url": _public_url(folder_path, fname), "path": filepath}


def _normalize_content_for_export(content: Any) -> List[dict]:
    """Normalize content for Word export."""
    from ..processing.block_utils import _convert_markdown_to_structured
    
    def normalize_item(item: Any) -> List[dict]:
        if item is None:
            return []
        if isinstance(item, str):
            structured = _convert_markdown_to_structured(item)
            normalized_blocks: List[dict] = []
            for block in structured:
                normalized_block = dict(block)
                if normalized_block.get("text"):
                    raw_text = normalized_block["text"]
                    plain, formatted = _normalize_markup_text(raw_text)
                    normalized_block["text"] = plain
                    normalized_block["formatted"] = formatted
                if normalized_block["type"] in {"paragraph", "text", "body", "description", "summary", "title"}:
                    normalized_block["segments"] = _parse_paragraph_segments(raw_text)
                normalized_blocks.extend(_finalize_normalized_block(normalized_block))
            return normalized_blocks
        if isinstance(item, dict):
            normalized: dict = {}
            item_type = (item.get("type") or "").strip().lower()

            if not item_type:
                if item.get("items"):
                    item_type = "list"
                elif item.get("cells") or item.get("rows") or item.get("table"):
                    item_type = "table"
                else:
                    item_type = "paragraph"

            normalized["type"] = item_type

            text_value = _extract_block_text(item)
            if item_type == "table":
                from ..processing.text_utils import _normalize_table_rows
                data = (
                    item.get("data")
                    or item.get("rows")
                    or item.get("cells")
                    or item.get("content")
                )
                normalized["data"] = _normalize_table_rows(data)
            else:
                raw_source = item.get("text") or item.get("title") or item.get("content") or ""
                normalized_plain, normalized_formatted = _normalize_markup_text(raw_source)
                if normalized_plain:
                    normalized["text"] = normalized_plain
                    normalized["formatted"] = normalized_formatted
                    normalized["raw_text"] = raw_source
                    if item_type in {"paragraph", "text", "body", "description", "summary", "title"}:
                        normalized["segments"] = _parse_paragraph_segments(raw_source)

            if item.get("children"):
                normalized["children"] = normalize_item(item.get("children"))

            if item_type == "list" and item.get("items"):
                normalized["items"] = []
                for list_entry in item.get("items") or []:
                    # Handle dict entries with "text" and "formatted" keys directly
                    if isinstance(list_entry, dict):
                        entry_text = list_entry.get("text")
                        entry_formatted = list_entry.get("formatted")
                        if entry_text:
                            entry_plain = str(entry_text).strip()
                            normalized["items"].append({"text": entry_plain, "formatted": entry_formatted})
                    else:
                        entry_plain, entry_formatted = _normalize_markup_text(list_entry)
                        if entry_plain:
                            normalized["items"].append({"text": entry_plain, "formatted": entry_formatted})
            return _finalize_normalized_block(normalized)
        if isinstance(item, list):
            return normalize_list(item)
        plain, formatted = _normalize_markup_text(item)
        if not plain:
            return []
        fallback_block = {"type": "paragraph", "text": plain, "formatted": formatted}
        fallback_block["segments"] = _parse_paragraph_segments(plain)
        return _finalize_normalized_block(fallback_block)

    def normalize_list(items: List[Any] | None) -> List[dict]:
        if not isinstance(items, list):
            return []
        result: List[dict] = []
        for child in items:
            result.extend(normalize_item(child))
        return result

    return normalize_item(content)


def _render_structured_docx(doc: Document, blocks: List[Any], depth: int = 1) -> None:
    """Render structured blocks to Word document."""
    for block in blocks:
        if isinstance(block, str):
            _add_docx_paragraph(doc, block)
            continue
        if not isinstance(block, dict):
            continue

        block_type = (block.get("type") or "").strip().lower()
        text = _extract_block_text(block)

        # Table handling
        if block_type == "table":
            from ..processing.text_utils import _normalize_table_rows
            table_data = _normalize_table_rows(
                block.get("data") or block.get("rows") or block.get("cells") or block.get("content") or []
            )
            if table_data:
                cols = max((len(row) for row in table_data), default=0)
                if cols > 0:
                    table = doc.add_table(rows=len(table_data), cols=cols)
                    try:
                        table.style = "Table Grid"
                    except KeyError:
                        pass

                    for i, row in enumerate(table_data):
                        for j in range(cols):
                            cell_text = str(row[j]) if j < len(row) else ""
                            cell = table.cell(i, j)
                            cell.text = cell_text

                            for para in cell.paragraphs:
                                para.paragraph_format.space_before = DocxPt(0)
                                para.paragraph_format.space_after = DocxPt(0)

                            if i == 0:
                                for para in cell.paragraphs:
                                    para.alignment = WD_ALIGN_PARAGRAPH.CENTER
                                    for run in para.runs:
                                        run.bold = True

                    doc.add_paragraph()
            continue

        # Title handling
        if block_type == "title":
            _add_docx_heading(doc, text, level=1, centered=True)
            if block.get("children"):
                _render_structured_docx(doc, block.get("children"), depth)
            continue

        # Section handling
        if block_type == "section":
            _add_docx_heading(doc, text or block.get("title") or "Section", level=min(depth + 1, 3))
            if block.get("children"):
                _render_structured_docx(doc, block.get("children"), depth + 1)
            continue

        # Paragraph-like handling
        if block_type in {"paragraph", "text", "body", "description", "summary"}:
            if text:
                _add_docx_paragraph(doc, text, block.get("formatted"))
            continue

        # Label paragraph handling
        if block_type == "label_paragraph":
            label_html = block.get("formatted_label") or block.get("label")
            body_html = block.get("formatted") or block.get("text")
            formatted_content = " ".join(part for part in (label_html, body_html) if part)
            plain_content = " ".join(
                part.strip() for part in (block.get("label"), block.get("text"))
                if part and str(part).strip()
            )
            _add_docx_paragraph(doc, plain_content, formatted_content or None)
            continue

        # Sources handling
        if block_type in {"sources", "source", "references"}:
            _add_docx_heading(doc, text or "Sources", level=min(depth + 1, 3))
            entries = block.get("children") or block.get("items") or []
            for entry in entries:
                entry_text = _extract_block_text(entry)
                entry_formatted = entry.get("formatted") if isinstance(entry, dict) else None
                if entry_text:
                    para = doc.add_paragraph()
                    _apply_formatted_html_to_paragraph(para, entry_formatted, entry_text)
                    para.paragraph_format.space_before = DocxPt(2)
                    para.paragraph_format.space_after = DocxPt(2)
                    para.paragraph_format.left_indent = DocxPt(12)
                    run = para.runs[0] if para.runs else para.add_run()
                    run.italic = True
            continue

        # List handling
        if "list" in block_type or "bullet" in block_type:
            entries = block.get("items") or block.get("children") or block.get("entries") or []
            if not entries and block_type in {"bullet", "list_item"}:
                text_value = block.get("text")
                if text_value:
                    entries = [text_value]
            for entry in entries:
                entry_text = _extract_block_text(entry)
                entry_formatted = entry.get("formatted") if isinstance(entry, dict) else None
                if not entry_text:
                    continue
                para = doc.add_paragraph()
                # Always add "• " prefix to ensure bullet character is in text content
                # when text is extracted from the DOCX file
                entry_text = "• " + entry_text
                # Also add "• " prefix to formatted content if present
                if entry_formatted:
                    entry_formatted = "• " + entry_formatted
                # Apply bullet list style if available, otherwise use normal style
                try:
                    para.style = doc.styles["List Bullet"]
                except KeyError:
                    para.style = doc.styles["Normal"]
                para.paragraph_format.left_indent = DocxPt(12 * depth)
                para.paragraph_format.space_before = DocxPt(2)
                para.paragraph_format.space_after = DocxPt(2)
                _apply_formatted_html_to_paragraph(para, entry_formatted, entry_text)
                if isinstance(entry, dict) and entry.get("children"):
                    _render_structured_docx(doc, entry.get("children"), depth + 1)
            continue

        # Unknown type handling
        if text:
            _add_docx_paragraph(doc, text)
        if block.get("children"):
            _render_structured_docx(doc, block.get("children"), depth + 1)


def _add_docx_paragraph(doc: Document, text: str, formatted: Optional[str] = None) -> None:
    """Add a paragraph to the document."""
    if not text and not formatted:
        return
    paragraph = doc.add_paragraph()
    _apply_formatted_html_to_paragraph(paragraph, formatted, text)
    try:
        paragraph.style = doc.styles['Normal']
    except KeyError:
        pass
    paragraph.paragraph_format.space_before = DocxPt(4)
    paragraph.paragraph_format.space_after = DocxPt(6)


def _apply_formatted_html_to_paragraph(paragraph, formatted_html: Optional[str], fallback_text: str) -> None:
    """Apply HTML formatting to a paragraph."""
    while paragraph.runs:
        paragraph.runs[0]._element.getparent().remove(paragraph.runs[0]._element)
    content = formatted_html or fallback_text or ""
    if not content:
        return
    from bs4 import BeautifulSoup
    soup = BeautifulSoup(content, "html.parser")

    def recurse(node, bold=False, italic=False):
        for child in node.children:
            if isinstance(child, str):
                if child:
                    run = paragraph.add_run(child)
                    run.bold = bold
                    run.italic = italic
            elif child.name in {"strong", "b"}:
                recurse(child, bold=True, italic=italic)
            elif child.name in {"em", "i"}:
                recurse(child, bold=bold, italic=True)
            elif child.name == "br":
                paragraph.add_run().add_break()
            else:
                recurse(child, bold=bold, italic=italic)

    recurse(soup)


def _add_docx_heading(doc: Document, text: str, level: int = 2, centered: bool = False) -> None:
    """Add a heading to the document."""
    if not text:
        return
    paragraph = doc.add_paragraph(text)
    style_name = f"Heading {min(max(level, 1), 3)}"
    try:
        paragraph.style = doc.styles[style_name]
    except KeyError:
        run = paragraph.runs[0] if paragraph.runs else paragraph.add_run()
        run.font.size = DocxPt(18 if level <= 2 else 14)
        run.font.bold = True
    if centered:
        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    paragraph.paragraph_format.space_before = DocxPt(6)
    paragraph.paragraph_format.space_after = DocxPt(6)


def _extract_block_text(block: Any) -> str:
    """Extract text from a block."""
    if isinstance(block, str):
        return block.strip()
    if not isinstance(block, dict):
        return ""
    block_type = (block.get("type") or "").strip().lower()
    if block_type == "table":
        return ""
    for key in ("text", "title", "description", "label", "name"):
        value = block.get(key)
        if value:
            return str(value).strip()
    return ""