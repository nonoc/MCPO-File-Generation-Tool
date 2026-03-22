"""Regression tests that trace content loss through the export pipeline."""

import asyncio
import base64
import os
import re
import shutil
import sys
from pathlib import Path
import zlib

import markdown2
import pytest
from bs4 import BeautifulSoup
from docx import Document
from reportlab.platypus import Paragraph as RLParagraph, Table as RLTable

TOOLS_PATH = Path(__file__).resolve().parent.parent / "LLM_Export" / "tools"
if str(TOOLS_PATH) not in sys.path:
    sys.path.insert(0, str(TOOLS_PATH))

import importlib  # noqa: E402

PDF_STREAM_RE = re.compile(rb"<<(.*?)>>\s*stream\s*(.*?)\s*endstream", re.S)

SECOND_PARAGRAPH_PHRASE = (
    "This is the second paragraph that should appear in the PDF body."
)
"""A phrase we track in every export to ensure the body is preserved."""

BULLET_PHRASE = "First bullet point in the list"
STRUCTURED_TITLE = "Structured Title"
STRUCTURED_SECTION = "Section Alpha"
STRUCTURED_CHILD_PARAGRAPH = "Paragraph inside Section Alpha."
STRUCTURED_DOCX_SECTION = "Section 7.1"
STRUCTURED_DOCX_BULLETS = [
    "Docx bullet alpha",
    "Docx bullet beta",
]
STRUCTURED_SOURCES = [
    "Source Doc A - Legal memo",
    "Source Doc B - Compliance checklist",
]
STRUCTURED_BULLET_SECTION = "Add structured bullet insights here"
STRUCTURED_BULLET_ONE = "First structured bullet item"
STRUCTURED_BULLET_TWO = "Second structured bullet item"
MULTI_SECTION_FIRST = "First Structured Section"
MULTI_SECTION_SECOND = "Second Structured Section"
MULTI_SECTION_FIRST_PARAGRAPH = "Paragraph for first structured section."
MULTI_SECTION_SECOND_PARAGRAPH = "Paragraph for second structured section."
SOURCES_SECTION_TITLE = "Sources Cited"
SOURCES_LIST = [
    "Source A: Background whitepaper",
    "Source B: Compliance document",
]
MARKDOWN_TABLE_REPORT = """# Table Report\n**Bold insight** is key.\n| Column | Value |\n| ------ | ----- |\n| Users  | 1,000 |\n"""

MIXED_PARAGRAPH_SAMPLE = """Context before labels.
**Impact:** Adoption velocity spikes with personalization.
**Mitigation:** Continue monitoring KPIs and keep rollback safeguards ready.
- Validate data flows nightly
- Monitor dashboards for anomalies
Closing statement keeps customers informed."""

MARKDOWN_BODY = f"""# Main Title

This is the first paragraph of the document. It provides context for
the diagnostic scenario.

{SECOND_PARAGRAPH_PHRASE}

## Section One

Here is a supporting paragraph for Section One.

- {BULLET_PHRASE}
- Second bullet detailing behavior
- Third bullet to close the series

### Subsection

Final paragraph wraps up the markdown story.
"""

INTEGRATION_REPORT_BODY = f"""# Executive Summary

The project continues to meet expectations and sets the stage for the
next phase.

{SECOND_PARAGRAPH_PHRASE}

## Key Findings

The analysis revealed several patterns worth highlighting.

- {BULLET_PHRASE}
- Data quality remains excellent across the sample.
- User engagement increased by 25% compared to the previous quarter.

## Recommendations

1. Continue the current data collection methodology.
2. Expand user engagement initiatives to additional regions.
3. Allocate additional resources to the successful feature teams.

## Conclusion

The report confirms that the investments are driving the desired outcomes.
"""


@pytest.fixture(scope="session")
def file_export_module(tmp_path_factory):
    os.environ.setdefault("FILE_EXPORT_BASE_URL", "http://localhost:9003/files")
    os.environ.setdefault("LOG_LEVEL", "DEBUG")
    export_dir = tmp_path_factory.mktemp("export_test")
    os.environ["FILE_EXPORT_DIR"] = str(export_dir)
    # Import from the tools package to handle relative imports correctly
    module = importlib.import_module("LLM_Export.tools.file_export_mcp")
    return module


def extract_pdf_text(pdf_path: Path) -> str:
    raw_bytes = pdf_path.read_bytes()
    decoded_chunks: list[bytes] = []
    for header, stream in PDF_STREAM_RE.findall(raw_bytes):
        chunk = stream.strip()
        if b"/ASCII85Decode" in header:
            try:
                chunk = base64.a85decode(chunk, adobe=True)
            except ValueError:
                continue
        if b"/FlateDecode" in header:
            try:
                chunk = zlib.decompress(chunk)
            except zlib.error:
                continue
        decoded_chunks.append(chunk)
    return b" ".join(decoded_chunks).decode("latin-1", errors="ignore")


def _assert_contains_second_paragraph_and_bullet(text: str) -> None:
    assert SECOND_PARAGRAPH_PHRASE in text
    assert BULLET_PHRASE in text


def _assert_contains_second_paragraph(text: str) -> None:
    assert SECOND_PARAGRAPH_PHRASE in text


def test_a_pdf_export_with_string_markdown(file_export_module):
    module = file_export_module
    folder = module._generate_unique_folder()
    result = module._create_pdf(MARKDOWN_BODY, "body_check.pdf", folder_path=folder)
    pdf_path = Path(result["path"])
    assert pdf_path.exists(), "PDF file does not exist"
    assert pdf_path.stat().st_size > 1800, "PDF file size is unexpectedly small"
    pdf_text = extract_pdf_text(pdf_path)
    _assert_contains_second_paragraph_and_bullet(pdf_text)
    structured_blocks = module._convert_markdown_to_structured(MARKDOWN_BODY)
    assert len(structured_blocks) > 3


def test_b_docx_export_with_string_markdown(file_export_module):
    module = file_export_module
    folder = module._generate_unique_folder()
    result = module._create_word(MARKDOWN_BODY, "body_check.docx", folder_path=folder)
    doc = Document(result["path"])
    paragraphs = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
    joined_text = "\n".join(paragraphs)
    _assert_contains_second_paragraph(joined_text)
    assert len(paragraphs) > 5, "DOCX export should produce multiple paragraphs"


def test_c_markdown_conversion_sanity(file_export_module):
    structured = file_export_module._convert_markdown_to_structured(MARKDOWN_BODY)
    texts = " ".join(block.get("text", "") for block in structured)
    _assert_contains_second_paragraph_and_bullet(texts)
    assert len(structured) > 4
    block_types = {block.get("type") for block in structured}
    assert "paragraph" in block_types
    assert "heading" in block_types or "title" in block_types
    assert "bullet" in block_types or any("list" in t for t in block_types if t)


def test_d_html_render_pipeline_sanity(file_export_module):
    html = markdown2.markdown(MARKDOWN_BODY, extras=["fenced-code-blocks", "tables"])
    soup = BeautifulSoup(html, "html.parser")
    assert len(soup.find_all(["h1", "h2", "h3"])) >= 2
    assert len(soup.find_all("p")) >= 3
    assert len(soup.find_all("li")) >= 3
    story = file_export_module.render_html_elements(soup)
    paragraph_texts = [
        elem.getPlainText()
        for elem in story
        if isinstance(elem, RLParagraph)
    ]
    assert len(story) > 6
    assert len(paragraph_texts) > 1
    _assert_contains_second_paragraph(" ".join(paragraph_texts))


def test_e_integration_create_file(file_export_module):
    module = file_export_module
    export_dir = Path(os.environ["FILE_EXPORT_DIR"])
    before = set(export_dir.iterdir())
    payload = {
        "format": "pdf",
        "filename": "integration_report.pdf",
        "title": "Quarterly Report",
        "content": INTEGRATION_REPORT_BODY,
    }
    result = asyncio.run(module.create_file(payload, persistent=True))
    after = set(export_dir.iterdir())
    new_dirs = [path for path in after - before if path.is_dir()]
    assert len(new_dirs) == 1
    folder = new_dirs[0]
    pdf_candidates = list(folder.glob("*.pdf"))
    assert pdf_candidates, "No PDF generated by create_file"
    pdf_path = pdf_candidates[0]
    assert pdf_path.stat().st_size >= 2000
    pdf_text = extract_pdf_text(pdf_path)
    _assert_contains_second_paragraph_and_bullet(pdf_text)
    assert result["url"].endswith("integration_report.pdf")
    shutil.rmtree(folder, ignore_errors=True)


def test_f_pdf_export_with_structured_sections(file_export_module):
    module = file_export_module
    structured_content = [
        {"type": "title", "text": STRUCTURED_TITLE},
        {
            "type": "section",
            "title": STRUCTURED_SECTION,
            "children": [
                {"type": "paragraph", "text": STRUCTURED_CHILD_PARAGRAPH},
                {"type": "paragraph", "text": "Another child paragraph."},
            ],
        },
    ]

    export_dir = Path(os.environ["FILE_EXPORT_DIR"])
    before = set(export_dir.iterdir())
    payload = {
        "format": "pdf",
        "filename": "structured_nested.pdf",
        "content": structured_content,
    }
    result = asyncio.run(module.create_file(payload, persistent=True))
    after = set(export_dir.iterdir())
    new_dirs = [path for path in after - before if path.is_dir()]
    assert new_dirs, "No new export folder created"
    folder = new_dirs[0]
    pdf_candidates = list(folder.glob("*.pdf"))
    assert pdf_candidates, "Structured PDF not generated"
    pdf_path = pdf_candidates[0]
    pdf_text = extract_pdf_text(pdf_path)
    assert STRUCTURED_TITLE in pdf_text
    assert STRUCTURED_SECTION in pdf_text, "Section heading missing from structured payload"
    assert STRUCTURED_CHILD_PARAGRAPH in pdf_text, "Paragraph child missing from structured payload"
    shutil.rmtree(folder, ignore_errors=True)


def test_g_structured_bullet_rendering(file_export_module):
    module = file_export_module
    structured_content = [
        {"type": "title", "text": "Bullet Report"},
        {
            "type": "section",
            "title": STRUCTURED_BULLET_SECTION,
            "children": [
                {"type": "paragraph", "text": "Bullet list below captures key actions."},
                {"type": "list", "items": [STRUCTURED_BULLET_ONE, STRUCTURED_BULLET_TWO]},
            ],
        },
    ]

    folder = Path(module._generate_unique_folder())
    result = module._create_pdf(structured_content, "structured_bullets.pdf", folder_path=str(folder))
    pdf_path = Path(result["path"])
    assert pdf_path.exists()
    pdf_text = extract_pdf_text(pdf_path)
    assert STRUCTURED_BULLET_SECTION in pdf_text
    assert STRUCTURED_BULLET_ONE in pdf_text
    assert STRUCTURED_BULLET_TWO in pdf_text
    shutil.rmtree(folder, ignore_errors=True)


def test_h_structured_multiple_sections(file_export_module):
    module = file_export_module
    structured_content = [
        {"type": "title", "text": "Multi-section Report"},
        {"type": "section", "title": MULTI_SECTION_FIRST, "children": [{"type": "paragraph", "text": MULTI_SECTION_FIRST_PARAGRAPH}]},
        {"type": "section", "title": MULTI_SECTION_SECOND, "children": [{"type": "paragraph", "text": MULTI_SECTION_SECOND_PARAGRAPH}]},
    ]

    folder = Path(module._generate_unique_folder())
    result = module._create_pdf(structured_content, "structured_sections.pdf", folder_path=str(folder))
    pdf_text = extract_pdf_text(Path(result["path"]))
    assert MULTI_SECTION_FIRST in pdf_text
    assert MULTI_SECTION_SECOND in pdf_text
    assert pdf_text.index(MULTI_SECTION_FIRST) < pdf_text.index(MULTI_SECTION_SECOND)
    shutil.rmtree(folder, ignore_errors=True)


def test_i_structured_sources_section(file_export_module):
    module = file_export_module
    structured_content = [
        {"type": "title", "text": "Sources Report"},
        {
            "type": "section",
            "title": "Context",
            "children": [
                {"type": "paragraph", "text": "Contextual paragraph."},
                {"type": "sources", "title": SOURCES_SECTION_TITLE, "children": [{"type": "paragraph", "text": item} for item in SOURCES_LIST]},
            ],
        },
    ]

    folder = Path(module._generate_unique_folder())
    result = module._create_pdf(structured_content, "structured_sources.pdf", folder_path=str(folder))
    pdf_text = extract_pdf_text(Path(result["path"]))
    assert SOURCES_SECTION_TITLE in pdf_text
    for source in SOURCES_LIST:
        assert source in pdf_text
    shutil.rmtree(folder, ignore_errors=True)


def test_j_docx_structured_rendering(file_export_module):
    module = file_export_module
    structured_content = [
        {"type": "title", "text": "DOCX Structured Report"},
        {
            "type": "section",
            "title": STRUCTURED_DOCX_SECTION,
            "children": [
                {"type": "paragraph", "text": "Paragraph within section to verify spacing."},
                {"type": "list", "items": STRUCTURED_DOCX_BULLETS},
            ],
        },
    ]

    folder = Path(module._generate_unique_folder())
    result = module._create_word(structured_content, "structured_docx.docx", folder_path=str(folder))
    doc = Document(result["path"])
    heading_texts = [p.text.strip() for p in doc.paragraphs if p.text.strip() and getattr(p.style, "name", "").startswith("Heading")]
    assert STRUCTURED_DOCX_SECTION in heading_texts
    bullet_paragraphs = [p for p in doc.paragraphs if any(bullet in p.text for bullet in STRUCTURED_DOCX_BULLETS)]
    assert bullet_paragraphs, "Bullets should render as separate paragraphs"
    assert all(getattr(p.style, "name", "") in ("List Bullet", "Normal") for p in bullet_paragraphs)
    assert STRUCTURED_DOCX_SECTION in " ".join(heading_texts)
    shutil.rmtree(folder, ignore_errors=True)


def test_k_docx_structured_sources(file_export_module):
    module = file_export_module
    structured_content = [
        {"type": "title", "text": "Sources Report"},
        {
            "type": "sources",
            "title": "Sources",
            "children": [{"type": "paragraph", "text": source} for source in STRUCTURED_SOURCES],
        },
    ]

    folder = Path(module._generate_unique_folder())
    result = module._create_word(structured_content, "structured_sources.docx", folder_path=str(folder))
    doc = Document(result["path"])
    source_paragraphs = [p for p in doc.paragraphs if any(source in p.text for source in STRUCTURED_SOURCES)]
    assert len(source_paragraphs) == len(STRUCTURED_SOURCES)
    assert all(any(run.italic for run in p.runs) for p in source_paragraphs)
    shutil.rmtree(folder, ignore_errors=True)


def test_l_markdown_table_and_bold_export(file_export_module):
    module = file_export_module
    folder = Path(module._generate_unique_folder())
    result_pdf = module._create_pdf(MARKDOWN_TABLE_REPORT, "table_report.pdf", folder_path=str(folder))
    pdf_text = extract_pdf_text(Path(result_pdf["path"]))
    assert "Column" in pdf_text
    assert "1,000" in pdf_text
    assert "|" not in pdf_text
    result_docx = module._create_word(MARKDOWN_TABLE_REPORT, "table_report.docx", folder_path=str(folder))
    doc = Document(result_docx["path"])
    assert doc.tables, "Table not rendered"
    table = doc.tables[0]
    assert table.cell(0, 0).text == "Column"
    assert table.cell(1, 1).text == "1,000"
    bold_para = next((p for p in doc.paragraphs if "Bold insight" in p.text), None)
    assert bold_para is not None
    bold_run = next((run for run in bold_para.runs if "Bold insight" in run.text), None)
    assert bold_run is not None and bold_run.bold
    shutil.rmtree(folder, ignore_errors=True)


def test_m_table_block_story_includes_table(file_export_module):
    module = file_export_module
    structured = [{"type": "table", "data": [["Column", "Value"], ["Users", "1,000"]]}]
    normalized = module.normalize_content_for_export(structured)
    story = module.render_structured_content(normalized)
    assert any(isinstance(elem, RLTable) for elem in story), "Table block missing from PDF story"


def test_n_mixed_paragraph_structuring(file_export_module):
    module = file_export_module
    folder_pdf = Path(module._generate_unique_folder())
    pdf_result = module._create_pdf(MIXED_PARAGRAPH_SAMPLE, "mixed_paragraph.pdf", folder_path=str(folder_pdf))
    pdf_text = extract_pdf_text(Path(pdf_result["path"]))
    assert pdf_text.count("Impact:") == 1
    assert pdf_text.count("Mitigation:") == 1
    assert "Validate data flows nightly" in pdf_text
    assert "Monitor dashboards for anomalies" in pdf_text
    shutil.rmtree(folder_pdf, ignore_errors=True)
    folder_docx = Path(module._generate_unique_folder())
    docx_result = module._create_word(MIXED_PARAGRAPH_SAMPLE, "mixed_paragraph.docx", folder_path=str(folder_docx))
    doc = Document(docx_result["path"])
    label_paragraph = next((p for p in doc.paragraphs if "Impact:" in p.text), None)
    assert label_paragraph is not None
    label_run = next((run for run in label_paragraph.runs if "Impact" in run.text), None)
    assert label_run is not None and label_run.bold
    bullet_texts = ["Validate data flows nightly", "Monitor dashboards for anomalies"]
    bullet_paragraphs = [p for p in doc.paragraphs if any(text in p.text for text in bullet_texts)]
    assert len(bullet_paragraphs) == len(bullet_texts)
    assert sum(1 for p in doc.paragraphs if "Mitigation:" in p.text) == 1
    shutil.rmtree(folder_docx, ignore_errors=True)
